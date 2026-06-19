# -*- coding: utf-8 -*-
"""Build the EduScheduler User Guide (Word .docx)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1E, 0x27, 0x61)
AMBER = RGBColor(0xC0, 0x73, 0x00)
SLATE = RGBColor(0x44, 0x4C, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for hs, sz, col in [('Heading 1', 18, NAVY), ('Heading 2', 14, NAVY), ('Heading 3', 12, AMBER)]:
    st = doc.styles[hs]
    st.font.name = 'Cambria'; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = True
    st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(4)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, bold=False, color=None, size=10.5, align=None, font='Calibri'):
    cell.text = ''
    p = cell.paragraphs[0]
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text); r.font.bold = bold; r.font.size = Pt(size); r.font.name = font
    if color: r.font.color.rgb = color


def callout(title, body, fill='EEF2FB', barcolor=None):
    """A simple shaded one-cell note box."""
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    cell = t.cell(0, 0); shade(cell, fill)
    cell.margin_top = Pt(4)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + '  '); r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = NAVY
    r2 = p.add_run(body); r2.font.size = Pt(10.5); r2.font.color.rgb = SLATE
    doc.add_paragraph()


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; shade(c, '1E2761')
        set_cell(c, h, bold=True, color=WHITE, size=10)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            set_cell(cells[j], str(val), size=10, bold=(j == 0 and len(headers) <= 3))
            if (len(t.rows) % 2) == 1:
                shade(cells[j], 'F4F6FC')
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    doc.add_paragraph()
    return t


def para(text, size=11, color=None, bold=False, italic=False, space=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            r = p.add_run(it[0] + ' '); r.font.bold = True
            r2 = p.add_run(it[1])
        else:
            p.add_run(it)


def steps(items):
    for it in items:
        p = doc.add_paragraph(style='List Number'); p.paragraph_format.space_after = Pt(3)
        p.add_run(it)


# ════════════════════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('EduScheduler'); r.font.size = Pt(40); r.font.bold = True; r.font.name = 'Cambria'; r.font.color.rgb = NAVY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Smart College Timetable Management System'); r.font.size = Pt(18); r.font.color.rgb = AMBER; r.font.name = 'Cambria'
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(30)
r = p.add_run('User Guide & Documentation'); r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = NAVY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A complete, non-technical manual for administrators, departments, professors and students')
r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = SLATE
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CONTENTS
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('Contents', level=1)
toc = [
    '1.  Introduction', '2.  Key Terms (Glossary)', '3.  Getting Started & Logging In',
    '4.  User Roles & What Each Can Do', '5.  Preparing Your Data — the Five Files',
    '6.  Uploading Your Data', '7.  Generating the Timetable',
    '8.  How the Timetable is Built (the Rules)', '9.  Viewing Timetables',
    '10.  Exporting & Sharing', '11.  Managing Accounts (Admin)',
    '12.  Adding & Editing Without CSV', '13.  Updating & Re-uploading',
    '14.  Troubleshooting & Tips', '15.  Appendix — CSV Column Reference',
]
for t in toc:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.size = Pt(11.5)
doc.add_page_break()

# ── 1. INTRODUCTION ─────────────────────────────────────────────────────────
doc.add_heading('1.  Introduction', level=1)
para('EduScheduler is a web application that automatically builds clash-free weekly timetables for a whole '
     'college. Instead of arranging classes by hand in spreadsheets, you describe your college in five simple '
     'files, upload them, and click one button. The system schedules every class for every group while '
     'respecting lunch breaks, lab requirements, professor workloads, fixed and unavailable times, and the '
     'fact that professors and rooms are often shared between departments.')
para('What you can do with it:', bold=True, space=2)
bullets([
    ('Generate', 'a complete timetable for all departments with a single click.'),
    ('View', 'timetables by section, professor, room or semester.'),
    ('Export', 'any timetable to PDF, spreadsheet (CSV) or a scannable QR code.'),
    ('Control access', 'so each person sees only what they are allowed to.'),
    ('Re-generate', 'instantly whenever the data changes.'),
])
callout('In plain terms:', 'You provide the facts (who teaches what, which rooms exist, which groups study which '
        'subjects); EduScheduler does the puzzle-solving.')

# ── 2. GLOSSARY ─────────────────────────────────────────────────────────────
doc.add_heading('2.  Key Terms (Glossary)', level=1)
table(['Term', 'Meaning'], [
    ['Department', 'An academic department, e.g. Computer Science & Engineering.'],
    ['Course', 'A programme within a department, e.g. B.TECH.'],
    ['Semester', 'The study term, numbered 1 to 8.'],
    ['Section', 'A class within a semester (e.g. I, II, III or A, B). Each section has one or more groups.'],
    ['Group', 'A sub-division of a section: G1, G2, G3, G4 (e.g. for splitting labs).'],
    ['Slot / Period', 'A one-hour teaching window (9:00–9:50, 9:50–10:40, …).'],
    ['Theory / Lab / Tutorial', 'Types of class. A lab takes two consecutive slots.'],
    ['NPTEL', 'An online course type, always scheduled after lunch.'],
    ['Elective', 'An optional subject. A semester’s electives run at the same time so a student picks one.'],
    ['Fixed time', 'A slot where a professor MUST teach a subject.'],
    ['Blocked time', 'A slot where a professor is unavailable (meeting, etc.).'],
], widths=[2.0, 4.7])

# ── 3. GETTING STARTED ──────────────────────────────────────────────────────
doc.add_heading('3.  Getting Started & Logging In', level=1)
para('Open the application in a web browser. You will be taken to the login screen, which has four tabs — one '
     'for each kind of user. Choose your tab, enter your details, and sign in.')
doc.add_heading('The default administrator', level=2)
para('On first use, an administrator account already exists:')
table(['Field', 'Value'], [['Username', 'admin'], ['Password', 'admin123']], widths=[2.0, 3.0])
callout('Important:', 'Change the administrator password after your first login (or create your own admin) — the '
        'default is only meant to get you started.')
doc.add_heading('How each role logs in', level=2)
bullets([
    ('Administrator / Department Admin —', 'enter the username and password created for you.'),
    ('Professor —', 'enter your name exactly as in the uploaded data, and your Teacher ID as the password.'),
    ('Student —', 'no password. Simply choose Department → Course → Semester → Section and click View.'),
])

# ── 4. ROLES ────────────────────────────────────────────────────────────────
doc.add_heading('4.  User Roles & What Each Can Do', level=1)
para('EduScheduler enforces access by role — a person can only reach the pages their role allows.')
table(['Capability', 'Admin', 'Dept Admin', 'Professor', 'Student'], [
    ['Upload data & generate timetable', 'All departments', 'Own department only', 'No', 'No'],
    ['Add / edit professors, sections, rooms', 'All', 'Own department only', 'No', 'No'],
    ['Create & manage login accounts', 'Yes', 'No', 'No', 'No'],
    ['View timetables', 'Any', 'Own department', 'Own schedule only', 'Own section only'],
    ['Export to PDF / CSV / QR', 'Yes', 'Yes', 'Own', 'Own'],
    ['Needs a password', 'Yes', 'Yes', 'Yes (Teacher ID)', 'No'],
], widths=[3.0, 1.2, 1.3, 1.2, 1.0])
callout('Department isolation:', 'A Department Admin cannot open, view or change another department’s data — even '
        'by typing the address directly. Each department’s timetable is independent.')

# ── 5. PREPARING DATA ───────────────────────────────────────────────────────
doc.add_heading('5.  Preparing Your Data — the Five Files', level=1)
para('Your college is described by five spreadsheet (CSV) files. On the CSV Import page you can download a '
     'ready-made template for each — it already contains example rows, so you only edit the values. Save each '
     'as a .csv file (Excel: File → Save As → CSV).')
para('The five files:', bold=True, space=2)
table(['File', 'Describes', 'Key information'], [
    ['Department Settings', 'Each department', 'Day start time and lunch start time.'],
    ['Professors', 'Each teacher', 'ID, workload limit, subjects, fixed & blocked times.'],
    ['Rooms', 'Classrooms & labs', 'Room code, capacity, which subjects a lab allows.'],
    ['Sections', 'Each class group', 'Semester, section, group, fixed room, holiday, size.'],
    ['Subjects', 'Each subject', 'Theory/lab/tutorial hours, type, which groups take it.'],
], widths=[1.7, 1.6, 3.4])
callout('Forgiving formats:', 'The importer understands common variations automatically — “3rd” or “Semester 3”, '
        '“MB-202” (room code) or “Room 202” (room name), upper or lower case. You don’t have to be exact.')
para('Full column-by-column details for every file are in the Appendix (Section 15).', italic=True, color=SLATE)

# ── 6. UPLOADING ────────────────────────────────────────────────────────────
doc.add_heading('6.  Uploading Your Data', level=1)
steps([
    'Sign in as Administrator (or Department Admin) and open the CSV Import page.',
    'Optionally download any template you need and fill it in.',
    'Select or drag each of your five CSV files into its matching box. You can upload some or all.',
    'Tick “Auto-generate timetable after import” if you want the timetable built immediately.',
    'Click Upload. The system reads the files, reports how many records were imported, and lists any warnings.',
])
callout('Warnings are normal:', 'A warning (e.g. “fixed room not found”, “Saturday is not a working day”) means '
        'one value was skipped — not that the upload failed. The rest of your data still imports.')
doc.add_heading('Department Admins', level=2)
para('When a Department Admin uploads, any professor or room that doesn’t name a department is automatically '
     'assigned to their department, and the upload only affects their own department.')

# ── 7. GENERATING ───────────────────────────────────────────────────────────
doc.add_heading('7.  Generating the Timetable', level=1)
para('Click Smart Generate (or tick auto-generate during upload). In seconds, the system places every class '
     'and shows the result. You can re-generate at any time — it always rebuilds from the current data.')
callout('Re-uploading replaces cleanly:', 'If a timetable already exists and you upload again, the old one is '
        'cleared and rebuilt from the new file. For a Department Admin, only their department is rebuilt; other '
        'departments keep their timetables.')

# ── 8. RULES ────────────────────────────────────────────────────────────────
doc.add_heading('8.  How the Timetable is Built (the Rules)', level=1)
para('You don’t configure these — the engine applies them all automatically. Understanding them helps you read '
     'the result and prepare good data.')
rules = [
    ('No overload', 'A professor is never given three classes — or three labs — back to back on one day.'),
    ('Fair workload', 'Teaching is spread across the week, and the least-loaded eligible professor is preferred.'),
    ('Lunch break', 'No classes are scheduled during a department’s lunch slot.'),
    ('NPTEL after lunch', 'Online (NPTEL) classes are always placed after the lunch break.'),
    ('No clashes', 'A room is never double-booked, and a professor (even one shared between departments) is never in two places at once.'),
    ('Two-hour labs', 'A lab occupies two consecutive slots in a lab room. While one group is in the lab, other groups run tutorials in the same two slots.'),
    ('Electives in parallel', 'All of a semester’s electives are placed in the same slot, in different rooms, so a student attends one.'),
    ('Fixed & blocked times', 'A professor’s “must-teach” slots are reserved; their “unavailable” slots are kept free.'),
    ('Smart rooms', 'A section’s fixed room is used if set; otherwise a free room in the same department is chosen.'),
    ('Holidays & start times', 'A section’s weekly free day is respected, and classes start at the section’s or department’s start time.'),
]
for t, d in rules:
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t + ' — '); r.font.bold = True; r.font.color.rgb = NAVY
    p.add_run(d)
callout('If something can’t fit:', 'When a class genuinely cannot be placed (e.g. too few rooms), the system '
        'shows a clear warning instead of forcing a clash. Adjust the data and re-generate.')

# ── 9. VIEWING ──────────────────────────────────────────────────────────────
doc.add_heading('9.  Viewing Timetables', level=1)
para('From the dashboard you can open any of these views:')
bullets([
    ('Section timetable —', 'one group’s full week.'),
    ('Combined section —', 'all groups of a section together.'),
    ('Semester timetable —', 'every section of a semester at once.'),
    ('Professor schedule —', 'a teacher’s personal week, with fixed (must-teach) and blocked times marked.'),
    ('Room schedule —', 'everything happening in a room across the week.'),
])
para('The dashboard’s left-hand tree lets you drill down: Department → Course → Semester → Section → Group.',
     italic=True, color=SLATE)

# ── 10. EXPORT ──────────────────────────────────────────────────────────────
doc.add_heading('10.  Exporting & Sharing', level=1)
table(['Format', 'Use it for'], [
    ['PDF', 'Printing or emailing a polished, ready-to-pin timetable.'],
    ['CSV (spreadsheet)', 'Opening the data in Excel or importing elsewhere.'],
    ['QR code', 'Putting on a notice board — students scan it to open the timetable on a phone.'],
], widths=[2.2, 4.5])
para('Every view (section, professor, room, semester) has its own export buttons.', italic=True, color=SLATE)

# ── 11. ACCOUNTS ────────────────────────────────────────────────────────────
doc.add_heading('11.  Managing Accounts (Admin)', level=1)
para('Only the Administrator can create accounts. Use the Add Department Admin button on the dashboard.')
steps([
    'Click Add Department Admin (opens the Manage Accounts page).',
    'Enter a username and a password, and choose the department this admin will manage.',
    'Click Create. The new Department Admin can now sign in and manage only that department.',
    'To remove an account, click Delete next to it. (The main Administrator cannot be deleted here.)',
])

# ── 12. MANUAL EDIT ─────────────────────────────────────────────────────────
doc.add_heading('12.  Adding & Editing Without CSV', level=1)
para('Besides bulk upload, you can add or change individual items from the dashboard buttons: '
     'Department, Course, Room, Section/Group, Subject and Professor.')
doc.add_heading('Adding a section or group', level=2)
steps([
    'Click “+ Group”.',
    'Choose the Course and Semester.',
    'Pick the Section from the dropdown — it lists the sections that already exist (e.g. I, II, III) plus A–E, or choose “Custom” to type a new name.',
    'Tick one or more Groups (G1–G4). One section record is created per group ticked.',
    'Optionally set a fixed room, a weekly free day, and a start time, then Save.',
])
callout('Tip:', 'To add a new group to an existing section, just pick that section’s name from the dropdown and '
        'tick the extra group — it’s added to the same section.')
doc.add_heading('Giving a professor a login', level=2)
para('A professor signs in with their name and a Teacher ID. IDs come from the Professors CSV, but you can also '
     'set or change one on the Edit Professor page (the “Teacher ID” field is the professor’s password).')

# ── 13. UPDATING ────────────────────────────────────────────────────────────
doc.add_heading('13.  Updating & Re-uploading', level=1)
bullets([
    ('Changed one file?', 'Re-upload it. The timetable rebuilds from the current data.'),
    ('Department-by-department —', 'a Department Admin’s re-upload only rebuilds their department; others are untouched.'),
    ('Shared professors —', 'when a different department uploads, a shared professor’s combined schedule updates automatically and stays clash-free.'),
])

# ── 14. TROUBLESHOOTING ─────────────────────────────────────────────────────
doc.add_heading('14.  Troubleshooting & Tips', level=1)
table(['Symptom', 'Likely cause & fix'], [
    ['A professor can’t log in', 'Name or Teacher ID doesn’t match the data. Check spelling, or set the ID on Edit Professor.'],
    ['“fixed room not found” warning', 'The section’s fixed room isn’t in the Rooms file. Add the room, or leave the fixed room blank to auto-assign.'],
    ['Some classes weren’t placed', 'Not enough rooms or slots for the rules. Add rooms/capacity and re-generate; read the warnings.'],
    ['A section shows the wrong groups', 'Each section name must be distinct. Use clear section names (I, II, III or A, B).'],
    ['A free day was ignored', 'Only Monday–Friday are supported as working/free days.'],
    ['Times look shifted', 'Use times on the slot grid (9:00, 9:50, 10:40 …). Off-grid times snap to the nearest slot.'],
], widths=[2.4, 4.3])

# ── 15. APPENDIX ────────────────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('15.  Appendix — CSV Column Reference', level=1)
para('Below is every column of each file. Columns marked “optional” may be left blank.', italic=True, color=SLATE)

doc.add_heading('Department Settings', level=2)
table(['Column', 'Meaning / example'], [
    ['Department', 'Department name, e.g. Computer Science & Engineering.'],
    ['Lunch_Start_time', 'When lunch begins, e.g. 12:20.'],
    ['Department_Start_time', 'When the day begins, e.g. 9:00.'],
], widths=[2.3, 4.4])

doc.add_heading('Professors', level=2)
table(['Column', 'Meaning / example'], [
    ['Department_name', 'The professor’s department.'],
    ['Teacher_id', 'Unique ID — also the professor’s login password, e.g. CS-AJS.'],
    ['Teacher_name', 'Full name — the professor’s login username.'],
    ['Max_Workload_Hours_per_week', 'Maximum teaching hours per week (state once per teacher).'],
    ['Subject Name', 'A subject this teacher takes (one per row; repeat the teacher on more rows).'],
    ['Program_name', 'Branch the teacher serves, e.g. CSE or “CSE,COE”.'],
    ['Course_name', 'Degree, e.g. B.TECH or M.TECH.'],
    ['Semester', 'Semester of the class, e.g. 6th.'],
    ['Section', 'Section name, e.g. CS-1.'],
    ['Group', 'G1, G2 or “G1,G2” for both. A different teacher may be named per group.'],
    ['Block_time_slot(day/time)', 'Unavailable time, e.g. “Tuesday,9:50 to 11:30” (optional). Use | for several.'],
    ['Fixed_time_slot(day/time)', 'Must-teach time, e.g. “Wednesday,9:50 to 11:30” (optional).'],
], widths=[2.7, 4.0])

doc.add_heading('Rooms', level=2)
table(['Column', 'Meaning / example'], [
    ['Department_name', 'Owning department.'],
    ['Room_id', 'Short code, e.g. MB-202 (sections refer to this).'],
    ['Room_name', 'Friendly name, e.g. Room 202.'],
    ['Room_type', 'classroom or lab.'],
    ['Capacity', 'Number of students it holds.'],
    ['Allowed_Subjects', '“all”, or the specific lab subject this lab hosts (e.g. OS Lab).'],
], widths=[2.3, 4.4])

doc.add_heading('Sections', level=2)
table(['Column', 'Meaning / example'], [
    ['Department', 'Department name.'],
    ['Semester', 'Semester 1–8 (accepts “3rd” etc.).'],
    ['section', 'Section name, e.g. I, II, III or A, B.'],
    ['group', 'G1, G2, G3 or G4.'],
    ['Fixed_room', 'Room code/name for theory (optional — blank = auto-assign).'],
    ['Course', 'Programme, e.g. B.TECH.'],
    ['Program Name', 'Display name (optional).'],
    ['Free_day', 'Weekly holiday, Monday–Friday (optional).'],
    ['Class_Count', 'Number of students (used to pick a big-enough room).'],
    ['Day / Section_Start_time', 'Per-day start time (optional). Blank = use department start time.'],
], widths=[2.3, 4.4])

doc.add_heading('Subjects', level=2)
table(['Column', 'Meaning / example'], [
    ['Department_name', 'Owning department.'],
    ['Subject_id', 'Subject code, e.g. AGCS-21301.'],
    ['Subject_Name', 'Subject name, e.g. DBMS.'],
    ['Sub_type', 'REGULAR, ELECTIVE or NPTEL.'],
    ['Theory_per_week_per_section', 'Theory periods per week (shared by both groups of the section). For NPTEL rows this is the weekly NPTEL session count.'],
    ['Tutorial_per_week_per_group', 'Tutorial sessions per week, per group.'],
    ['Lab_per_week_per_group', 'Lab sessions per week, per group (each is a 2-slot block; e.g. 1, or 2 for a major project).'],
    ['Allowed_groups', 'Which groups take it, e.g. “G1 G2” or a single “G1”.'],
    ['Course_name', 'Degree the subject belongs to, e.g. B.TECH.'],
    ['Program_name', 'Branch, e.g. CSE.'],
    ['Semester', 'Semester of the class, e.g. 4th.'],
], widths=[2.7, 4.0])

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— End of User Guide —'); r.font.italic = True; r.font.color.rgb = SLATE; r.font.size = Pt(10)

out = 'D:/college_timetable_v2_work/_deliverables/EduScheduler_User_Guide.docx'
doc.save(out)
print('Saved', out)
